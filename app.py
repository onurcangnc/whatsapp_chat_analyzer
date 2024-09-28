from flask import Flask, request, send_file, render_template, session, redirect, url_for, flash, after_this_request
import os
import pandas as pd
import re
from datetime import datetime
import tempfile
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as ExcelImage
from collections import defaultdict
import openai
from openai import RateLimitError, APIError
import time
from dotenv import load_dotenv
from docx import Document
import sys


app = Flask(__name__)

sys.stdout.reconfigure(encoding='utf-8')

# Load environment variables from .env file
load_dotenv()

app.secret_key = os.getenv('FLASK_SECRET_KEY', 'fallback_secret_key')

# Set the OpenAI API key from the environment variable
openai.api_key = os.getenv("OPENAI_API_KEY")

if not openai.api_key:
    raise ValueError("OpenAI API key not set. Please configure it in the environment variable.")


def parse_whatsapp_chat(file):
    message_pattern = r'\[(\d{1,2})\.(\d{1,2})\.(\d{4}),? (\d{2}:\d{2}:\d{2})\] ([^:]+): (.*)'
    chat_data = []
    current_message = None

    # Read and decode the file
    for line in file:
        line = line.decode('utf-8').strip()
        if not line:
            continue

        match = re.match(message_pattern, line)
        if match:
            day = match.group(1)
            month = match.group(2)
            year = match.group(3)
            time_str = match.group(4)
            sender = match.group(5)
            message = match.group(6)

            date_time_str = f'{day}.{month}.{year} {time_str}'
            try:
                date_time_obj = datetime.strptime(date_time_str, '%d.%m.%Y %H:%M:%S')
            except ValueError as ve:
                print(f"Date parsing error: {ve} for line: {line}")
                continue

            chat_data.append({
                'Date': date_time_obj.date(),
                'Time': date_time_obj.time(),
                'Datetime': date_time_obj,
                'Sender': sender,
                'Message': message
            })
            current_message = len(chat_data) - 1
        elif current_message is not None:
            # Handle multi-line messages
            chat_data[current_message]['Message'] += '\n' + line

    return pd.DataFrame(chat_data)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('fileup')
        if file and file.filename.endswith('.txt'):
            # Parse the uploaded file
            parsed_chat = parse_whatsapp_chat(file)

            # Check if parsed_chat is empty
            if parsed_chat.empty:
                flash("Parsed chat is empty. Please check the chat format and try again.", 'error')
                return redirect(url_for('index'))

            # Calculate average reply times for each sender
            avg_reply_times = defaultdict(list)
            previous_sender = None
            previous_time = None

            for i, row in parsed_chat.iterrows():
                current_sender = row['Sender']
                current_time = row['Datetime']

                if previous_sender and previous_sender != current_sender:
                    time_difference = (current_time - previous_time).total_seconds() / 60  # Convert to minutes
                    avg_reply_times[previous_sender].append(time_difference)

                previous_sender = current_sender
                previous_time = current_time

            avg_reply_times = {sender: sum(times) / len(times) for sender, times in avg_reply_times.items() if times}

            # Create a temporary Excel file
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx').name
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                parsed_chat.to_excel(writer, sheet_name='Parsed Chat', index=False)
                message_summary = parsed_chat['Sender'].value_counts().reset_index()
                message_summary.columns = ['Sender', 'Total Messages']
                message_summary.to_excel(writer, sheet_name='Message Summary', index=False)
                avg_reply_df = pd.DataFrame(list(avg_reply_times.items()), columns=['Sender', 'Average Reply Time (mins)'])
                avg_reply_df.to_excel(writer, sheet_name='Avg Reply Time', index=False)

            # Store the file path for download
            session['output_file'] = output_path
            session['file_downloaded'] = False  # Track download status

            flash("File processed successfully.", "success")
            return redirect(url_for('index'))

    return render_template('index.html')

@app.route('/download_report', methods=['GET'])
def download_report():
    report_file = session.get('report_file')
    if not report_file or not os.path.exists(report_file):
        flash("No report available to download.", "error")
        return redirect(url_for('index'))

    # Delete the report file after sending it to the user
    @after_this_request
    def remove_report_file(response):
        try:
            os.remove(report_file)
            session.pop('report_file', None)  # Clear the file from the session
        except Exception as e:
            print(f"Error deleting file: {str(e).encode('utf-8', errors='ignore')}")
        return response

    return send_file(report_file, as_attachment=True, download_name="psychological_report.txt")

@app.route('/download_psychological_report', methods=['GET'])
def download_psychological_report():
    report_file = session.get('report_file')
    if not report_file or not os.path.exists(report_file):
        flash("No psychological report available to download.", "error")
        return redirect(url_for('index'))

    # Delete the report file after sending it to the user
    @after_this_request
    def remove_report_file(response):
        try:
            os.remove(report_file)
            session.pop('report_file', None)  # Clear the file from the session
        except Exception as e:
            print(f"Error deleting file: {str(e)}")
        return response

    return send_file(report_file, as_attachment=True, download_name="psychological_analysis_report.docx")

@app.route('/download', methods=['GET'])
def download():
    output_file = session.get('output_file')
    if not output_file or not os.path.exists(output_file):
        flash("No file available to download.", "error")
        return redirect(url_for('index'))

    if session.get('file_downloaded'):
        flash("The file has already been downloaded.", "error")
        return redirect(url_for('index'))

    session['file_downloaded'] = True

    @after_this_request
    def remove_file(response):
        try:
            os.remove(output_file)
            session.pop('output_file', None)  # Clear session after file is removed
            session.pop('file_downloaded', None)  # Reset download state
            flash("Download complete. The session has been reset.", "success")
        except Exception as e:
            print(f"Error deleting file: {str(e).encode('utf-8', errors='ignore')}")
        return response

    return send_file(output_file, as_attachment=True, download_name="parsed_chat_with_reply_times.xlsx")

# Route for processing CSV/Excel files and generating the psychological report
@app.route('/upload_csv', methods=['POST'])
def upload_csv():
    file = request.files.get('file_csv')
    if file and (file.filename.endswith('.csv') or file.filename.endswith('.xlsx')):
        # Load the uploaded CSV or Excel file into a DataFrame
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)

        # Check if the file is empty
        if df.empty:
            flash("Uploaded file is empty. Please check the file and try again.", 'error')
            return redirect(url_for('index'))

        # Generate the psychological analysis report
        psychological_analysis = generate_psychological_report(df)

        # Generate a Word document with the analysis
        report_path = generate_word_report(psychological_analysis, "psychological_report.docx")

        # Store the report path for download
        session['report_file'] = report_path

        flash("Psychological analysis report generated successfully.", "success")
        return redirect(url_for('index'))

    flash("Invalid file format. Please upload a CSV or Excel file.", 'error')
    return redirect(url_for('index'))

# Function to generate a psychological report
def generate_psychological_report(chat_data):
    # Convert the chat data to a readable format
    chat_text = ""
    for index, row in chat_data.iterrows():
        chat_text += f"[{row['Date']} {row['Time']}] {row['Sender']}: {row['Message']}\n"

    # Prepare the prompt with chat data and the provided prompt description
    prompt = f"""
    For the duration of this conversation, act as an Industrial/Organizational Psychology expert 
    with a specialization in Human Resource Management and Management. Your task is to analyze 
    the following WhatsApp chat data and generate a psychological analysis report, considering 
    factors like emotions, relationships, psychological conditions, and communication patterns.

    WhatsApp Chat Data:
    {chat_text}
    
    Based on this data, please provide a comprehensive analysis of the individuals' 
    psychological states, emotions, and relationships.
    """

    # Retry logic for handling rate limits
    retry_attempts = 3  # Number of times to retry
    retry_delay = 20  # Seconds to wait before retrying

    for attempt in range(retry_attempts):
        try:
            # Send request to OpenAI API
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0
            )
            return response.choices[0].message.content
        
        except RateLimitError:
            print(f"Rate limit exceeded. Retrying in {retry_delay} seconds... (Attempt {attempt + 1}/{retry_attempts})")
            time.sleep(retry_delay)
        except APIError as e:
            return f"API error: {e}"

    return "Error: Rate limit exceeded after multiple attempts. Please try again later."



def generate_word_report(analysis_text, file_name):
    # Function to create a Word document with the analysis
    from docx import Document
    doc = Document()
    doc.add_heading('Psychological Analysis Report', 0)

    doc.add_paragraph(analysis_text)

    # Save the document as a temporary file
    word_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(word_file.name)

    return word_file.name

if __name__ == '__main__':
    app.run(debug=True)
